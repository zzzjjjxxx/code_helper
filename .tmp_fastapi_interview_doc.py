from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "docs" / "FastAPI面试问答_CodeHelper实战版.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
FONT = "Microsoft YaHei"


def set_font(run, size=10.5, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, GRAY)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, GRAY)


def add_body(doc, text, *, bold_label=None, italic=False, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_label:
        label = paragraph.add_run(bold_label)
        set_font(label, 10.5, DARK_BLUE, True)
    run = paragraph.add_run(text)
    set_font(run, 10.5, color or RGBColor(0, 0, 0), None, italic)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_font(run, 9, RGBColor(45, 45, 45))
    run.font.name = "Consolas"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    return paragraph


def add_q(doc, number, question, answer, project, followup):
    heading = doc.add_heading(f"Q{number}. {question}", level=2)
    heading.paragraph_format.keep_with_next = True
    add_body(doc, answer, bold_label="标准回答：")
    add_body(doc, project, bold_label="结合 code_helper：")
    add_body(doc, followup, bold_label="面试追问：", color=GRAY)


def q(question, answer, project, followup):
    return (question, answer, project, followup)


SECTIONS = [
    ("一、FastAPI 与 Web 基础", [
        q("FastAPI 是什么？", "FastAPI 是基于 ASGI 的现代 Python Web 框架。它用 Starlette 提供路由、中间件和异步能力，用 Pydantic 做请求校验与响应序列化，通常由 Uvicorn 运行。它的核心价值是类型标注、自动文档和异步支持。", "项目入口 apps/api/main.py 创建 FastAPI 应用，注册 tasks、workspace、events 和 security 路由；共享模型位于 packages/shared。", "继续追问时说明 FastAPI 本身不是服务器，Uvicorn 才是运行 ASGI 应用的服务器。"),
        q("FastAPI 与 Flask、Django 有什么区别？", "Flask 更轻量、约束更少；Django 是功能完整的全栈框架，包含 ORM、模板和后台；FastAPI 更强调类型校验、自动 OpenAPI 文档和 ASGI 异步能力。选择取决于项目规模、团队习惯和生态需求。", "code_helper 需要轻量 API、异步任务、SSE 和清晰的数据模型，因此采用 FastAPI，而前端独立使用 React。", "不要简单说 FastAPI 一定比 Flask 或 Django 快，应结合 I/O、数据库、部署和业务复杂度比较。"),
        q("FastAPI 为什么性能较好？", "主要原因不是 FastAPI 自动把所有代码变快，而是它基于 ASGI，能够高效处理并发 I/O；同时 Starlette 的路由和 Pydantic 的数据处理开销较低。真正的性能还取决于数据库、外部 API、阻塞代码和部署方式。", "项目中的 LLM 请求、SSE 事件流和后台执行都属于 I/O 场景，但文件读写和 pytest 仍可能阻塞，因此代码使用 asyncio.to_thread 或后台任务隔离。", "面试时主动指出基准测试、连接池、缓存和数据库索引通常比框架选择更影响实际性能。"),
        q("ASGI 和 WSGI 有什么区别？", "WSGI 是传统同步 Python Web 服务器接口，适合同步请求；ASGI 是面向异步和长连接的接口，支持 async/await、SSE 和 WebSocket。FastAPI 使用 ASGI，因此能处理异步 I/O 和流式响应。", "events.py 使用 StreamingResponse 推送 text/event-stream；这种长连接场景是 ASGI 相比 WSGI 的优势之一。", "如果使用同步数据库或阻塞函数，ASGI 仍可能被阻塞，不能把 ASGI 等同于全自动异步。"),
        q("Uvicorn 是做什么的？", "Uvicorn 是 ASGI 服务器，负责监听端口、接收 HTTP 请求、调用 FastAPI 应用并把响应写回客户端。FastAPI 负责应用逻辑，Uvicorn 负责运行时网络服务。", "开发时通常由 Uvicorn 运行 apps.api.main:app，前端 Vite 默认运行在 5173，后端默认监听 8000。", "生产环境还要考虑 Worker 数量、反向代理、超时和日志配置。"),
        q("Starlette 和 Pydantic 分别负责什么？", "Starlette 提供 ASGI、路由、中间件、请求响应和流式响应等 Web 能力；Pydantic 提供数据模型、类型转换、校验和序列化。FastAPI 将二者组合起来。", "TaskCreateRequest、TaskDetail、TaskEventModel 和 TestOutcome 是 Pydantic 模型；SSE 路由和 CORS 则依赖 Starlette/FastAPI 的 Web 能力。", "可以进一步解释 response_model 如何使用 Pydantic 做输出过滤和序列化。"),
        q("路径参数、查询参数和请求体有什么区别？", "路径参数标识资源，例如 /tasks/{task_id}；查询参数用于筛选、分页或控制行为，例如 ?after=10；请求体承载 JSON 等结构化数据，通常用于 POST 或 PUT。", "events.py 的 after 是查询参数，tasks.py 的 task_id 是路径参数，创建任务和聊天消息分别使用 TaskCreateRequest、TaskChatRequest 请求体。", "还要注意请求体通常不能用 GET 传递；路径参数应该用于资源定位，而不是塞入复杂 JSON。"),
        q("response_model 有什么作用？", "response_model 用于校验和序列化接口输出，也可以过滤不应该暴露的字段，并让 OpenAPI 文档准确描述响应结构。它与请求模型可以不同。", "tasks.py 为创建、查询、运行、回滚和聊天接口声明了响应模型，保证前端拿到的 JSON 字段稳定。", "如果返回的数据不符合模型，FastAPI 会在服务端暴露错误，应该通过测试尽早发现。"),
        q("如何统一处理异常？", "可以在路由层把已知业务异常转换成 HTTPException，也可以使用 app.exception_handler 注册全局处理器。未知异常应记录日志、返回通用错误，不向客户端泄露堆栈和密钥。", "tasks.py 将 KeyError 转成 404，将 RuntimeError 和 ValueError 转成 400 或 409；任务执行异常会更新任务状态并记录 task.failed 事件。", "继续追问异常分类：参数错误、资源不存在、权限错误、冲突和服务端错误应使用不同状态码。"),
    ]),
    ("二、异步、后台任务与生命周期", [
        q("async def 和普通 def 路由有什么区别？", "async def 路由可以在 await I/O 等待时把执行权交给事件循环；普通 def 路由通常由框架放到线程池执行。选择不是看语法偏好，而是看内部调用是否异步。", "run_task 是 async 接口，因为它要启动后台任务和等待异步状态；SQLite 的同步操作通过 asyncio.to_thread 避免直接阻塞事件循环。", "如果 async 路由内部调用同步的长时间函数，仍然会阻塞；如果整个调用链都是同步的，普通 def 反而更直观。"),
        q("什么是事件循环？", "事件循环负责调度协程。当协程执行 await 等待网络、文件或其他异步操作时，事件循环可以运行其他任务。它不是多核并行，而是协作式并发。", "SSE 连接、后台 Agent 任务和多个请求由 asyncio 调度；事件服务还为每个任务维护异步队列。", "要区分并发和并行：CPU 密集任务通常需要进程、线程池或独立任务系统。"),
        q("为什么不能在异步路由里直接执行阻塞代码？", "阻塞代码会占住事件循环线程，使其他请求、SSE 推送和定时状态刷新无法及时执行。常见阻塞代码包括同步数据库、subprocess、复杂文件扫描和同步 HTTP 请求。", "Agent 会读取工作区、运行 pytest、调用 LLM，这些操作可能阻塞，因此部分同步函数通过 asyncio.to_thread 执行。", "面试时可以说明：异步接口的性能取决于调用链是否真正非阻塞。"),
        q("asyncio.to_thread 解决什么问题？", "它把同步函数交给线程池运行，让当前协程在等待时不阻塞事件循环。适合文件 I/O、同步 SDK 或轻量阻塞操作，但不适合无限制地提交大量 CPU 密集任务。", "TaskService 使用 asyncio.to_thread 调用 SQLiteStore、search_memory、create_memory 等同步方法；workflow 也用它运行 pytest 和同步 LLM 调用。", "要注意线程安全、线程池大小、取消语义和数据库连接是否允许跨线程。"),
        q("asyncio.create_task 有什么风险？", "它会把协程放到后台执行，但任务异常、取消、重复启动和进程重启都需要额外管理。只保存 task 对象而不持久化状态，服务重启后任务可能丢失。", "TaskService 用 _running_tasks 和 _pending_follow_up_tasks 管理同一任务的执行与追问排队，并在完成后从字典移除。", "生产环境可使用 Redis、Celery、RQ 或其他持久化队列，数据库保存任务状态，进程内字典只做短期缓存。"),
        q("多 Worker 部署时后台任务有什么问题？", "多个 Worker 是多个独立进程，进程内字典、队列和 asyncio 任务互不共享。请求可能被不同 Worker 处理，导致状态不一致或同一任务重复运行。", "当前实现适合单机 MVP；如果部署多个 Worker，应把任务锁、队列和事件广播迁移到 Redis 或数据库，并保证任务幂等。", "这是一道很常见的生产化追问，回答限制和改进方案比只说“加锁”更重要。"),
        q("FastAPI 的生命周期如何管理？", "应用启动时可以初始化数据库、连接池、模型或外部客户端；关闭时释放连接和资源。新版项目通常使用 lifespan 上下文管理器统一处理。", "main.py 的 startup 事件创建 data、snapshots、workspace 目录，初始化 SQLiteStore、EventService、RollbackService 和 TaskService，并放入 app.state.container。", "可以追问如何改为 lifespan，以及测试时如何确保启动和关闭事件被执行。"),
    ]),
    ("三、依赖注入 Depends", [
        q("Depends 是什么？", "Depends 是 FastAPI 的依赖注入机制。路由声明自己需要什么对象，FastAPI 负责调用依赖函数、解析依赖参数并把结果传入路由。它能减少重复代码并方便测试替换。", "tasks.py 通过 Depends(get_container) 获取应用容器，路由不需要自己创建数据库或 TaskService。", "继续解释依赖可以嵌套，例如当前用户依赖 Token，Token 依赖请求头。"),
        q("如何注入数据库会话？", "通常定义一个依赖函数，在请求范围内创建 Session，使用 yield 返回，finally 中关闭；事务边界由服务层或依赖统一管理。异步数据库则使用 AsyncSession。", "当前项目的 SQLiteStore 由应用启动时创建并放入 AppContainer，路由通过容器访问；如果迁移到 SQLAlchemy，可以把 Session 依赖放在 deps.py。", "注意不要把全局共享的可变 Session 直接注入所有请求。"),
        q("如何用 Depends 实现登录验证？", "定义 get_current_user 依赖，从 Authorization 头读取 Bearer Token，验证签名和过期时间，再查询用户；需要登录的路由声明 user=Depends(get_current_user)。", "当前 code_helper 没有登录体系，但 get_container 已经展示了同样的依赖注入模式。未来可以在 deps.py 增加用户和权限依赖。", "要说明 Token 验证失败应返回 401，并附带适当的 WWW-Authenticate 信息。"),
        q("依赖可以嵌套吗？", "可以。一个依赖可以声明另一个依赖，FastAPI 会构建依赖图并按顺序解析。这样可以把认证、权限、租户和数据库会话组合起来。", "可以把 code_helper 的容器依赖扩展为 settings → store → task_service，路由只依赖最上层容器。", "嵌套太深会降低可读性，复杂业务仍应放入 Service，而不是全部堆在依赖函数里。"),
        q("依赖是否会缓存？", "同一个请求中，默认情况下相同依赖会复用结果，避免重复创建资源；可以使用 use_cache=False 关闭。依赖的缓存只在一次请求范围内，不是全局缓存。", "当前 get_container 返回请求所需的应用容器；数据库和事件服务本身在应用启动阶段创建，而不是每次路由调用都创建。", "面试时区分请求级缓存、进程级单例和分布式缓存。"),
        q("如何覆盖依赖做测试？", "可以使用 app.dependency_overrides[dependency] = fake_dependency，在测试结束后清理。这样可以替换真实数据库、登录用户或外部服务。", "code_helper 的测试通过 TestClient 和应用容器测试真实任务流程；如果增加认证或外部 LLM，可以覆盖相应依赖注入。", "要保证 fixture 结束后恢复 overrides，避免测试之间相互污染。"),
    ]),
    ("四、Pydantic 与数据建模", [
        q("请求模型和数据库模型有什么区别？", "请求模型描述外部 API 允许客户端提交的字段；数据库模型描述持久化结构；响应模型描述允许返回的字段。三者不应为了省事完全共用，否则容易暴露内部字段或让输入过度授权。", "TaskCreateRequest、TaskDetail、TaskEventModel 和 TaskSummary 作为 API 数据契约；SQLite 表结构则在 storage/sqlite.py 的 SCHEMA 中定义。", "面试时强调输入和输出是边界，数据库模型是内部实现，应该解耦。"),
        q("如何设置默认值和可选字段？", "必填字段不设置默认值；可选字段使用类型联合和 None；列表、字典等可变默认值使用 Field(default_factory=list)，避免多个实例共享同一个对象。", "TaskDetail 中 focus_paths、latest_retrieval、subgoals、memory 和 events 都使用 default_factory，保证每个响应对象拥有独立列表。", "Pydantic v2 中常用 model_dump、model_validate 等 API，注意不要混用旧版 dict。"),
        q("如何做自定义校验？", "可以使用 field_validator 或 model_validator 对单字段或多个字段做校验。校验逻辑应该描述输入约束，复杂业务规则应留给 Service 层。", "工作区路径是否位于允许范围、补丁 old 是否能匹配文件内容属于业务安全校验，不适合只放在 Pydantic 中，因为它依赖真实文件系统。", "区分格式校验和业务校验，避免在模型校验器里执行昂贵的 I/O。"),
        q("model_dump() 是什么？", "model_dump() 将 Pydantic 模型转换为 Python 字典，便于写入 JSON、数据库或作为 LLM 上下文。需要 JSON 兼容格式时可以使用 mode='json'。", "事件 payload、LLM 状态、测试报告和任务详情都调用 model_dump(mode='json')，确保 datetime、枚举等可以序列化。", "要知道 model_dump_json() 和 model_dump(mode='json') 的使用区别。"),
        q("Pydantic 如何处理嵌套数据？", "如果字段类型是另一个 BaseModel 或模型列表，Pydantic 会递归校验和序列化嵌套结构。输入错误时会返回结构化验证错误。", "TaskDetail 嵌套了 TestOutcome、RetrievalHit、TaskSubgoalRecord、MemoryRecord、SnapshotRecord 和 TaskEventModel。", "可以追问如何限制嵌套列表大小、字段长度和递归深度。"),
        q("为什么应该设置响应模型？", "响应模型能明确 API 契约、过滤内部字段、自动生成文档，并在开发阶段及时发现返回结构变化。它也是前端 TypeScript 类型的后端来源之一。", "前端 types.ts 基本镜像 shared models，TaskDetail 和 TaskChatResponse 的字段变化会影响前端渲染。", "响应模型不是权限控制的唯一手段，敏感数据仍应在服务层主动排除。"),
    ]),
    ("五、数据库、事务与存储", [
        q("如何管理数据库连接？", "数据库连接应集中管理，避免每个路由手动打开连接。应用启动时创建连接池或 Store，请求/服务层获取连接，操作完成后提交或关闭。", "当前项目用 SQLiteStore 统一连接和 schema 初始化，main.py 启动时调用 initialize；TaskService 不在每个路由中创建数据库。", "迁移 PostgreSQL 时可把 SQLiteStore 替换为 SQLAlchemy Session/AsyncSession，而不改变路由契约。"),
        q("什么是连接池？", "连接池预先维护一组数据库连接，请求到来时借用，完成后归还，减少频繁建立 TCP 和认证连接的成本。连接池大小要结合数据库承载能力和 Worker 数量设置。", "SQLite 不像 PostgreSQL 那样依赖传统连接池；当前 Store 每次连接数据库，但写入通过锁保护。生产迁移 PostgreSQL 后应配置连接池。", "连接池过大可能耗尽数据库连接，过小会导致请求排队。"),
        q("什么是事务？", "事务把多个数据库操作组织成一个不可分割的工作单元，遵循原子性、一致性、隔离性和持久性。失败时回滚，成功时提交。", "创建任务、更新任务状态、保存事件和记忆都需要考虑一致性；当前 SQLiteStore 使用连接上下文和写连接封装写入。", "面试时可举例说明：补丁文件写入成功但数据库更新失败时，如何记录恢复任务。"),
        q("如何处理提交和回滚？", "成功路径 commit，异常路径 rollback，最后关闭连接。事务范围应尽量短，避免把外部 LLM 请求或长时间测试包在数据库事务中。", "Agent 执行和 pytest 不应占用数据库事务；任务状态和事件在操作节点完成后分别持久化。文件安全回滚由快照系统负责，不等同于数据库 rollback。", "说明数据库回滚不能自动恢复文件系统，所以需要快照或版本控制。"),
        q("SQLAlchemy Session 是否线程安全？", "通常一个 Session 不能在多个线程或请求之间共享。应按请求或工作单元创建 Session，并让线程、异步上下文和事务边界清晰。", "当前 SQLiteStore 自己使用 threading.Lock 保护写入；如果换成 SQLAlchemy，应通过依赖为每个请求提供独立 Session。", "异步 AsyncSession 也不能随意跨并发任务共享。"),
        q("SQLite 和 PostgreSQL 有什么区别？", "SQLite 是嵌入式单文件数据库，部署简单，适合 MVP、单机和低并发；PostgreSQL 是独立数据库服务器，支持更强的并发、权限、索引、扩展和运维能力。", "code_helper 当前默认使用 data/app.db，便于本地开发；多 Worker、多人使用或任务量大时应迁移 PostgreSQL。", "迁移时要检查锁语义、日期类型、JSON、分页、索引和连接池，而不是只改连接字符串。"),
        q("Alembic 用来做什么？", "Alembic 是 SQLAlchemy 生态的数据库迁移工具，用版本化脚本描述表结构增删改，使开发、测试和生产环境可以一致升级或回滚 schema。", "当前项目通过 SQLiteStore.initialize 执行 CREATE TABLE IF NOT EXISTS，适合 MVP；正式项目应使用 Alembic 管理迁移。", "面试时指出 CREATE TABLE IF NOT EXISTS 不能替代完整迁移系统。"),
        q("如何避免 N+1 查询？", "N+1 是先查 N 个父记录，再为每个父记录单独查询子记录。可以使用 join、selectinload、批量查询或预取，减少数据库往返。", "任务详情会读取事件、快照、记忆、子目标和产物；如果任务量扩大，应批量加载关联数据并为 task_id、created_at 建索引。", "先用日志或 profiling 证明 N+1，再选择 eager loading，避免盲目优化。"),
    ]),
    ("六、安全、文件与命令执行", [
        q("JWT 登录流程是什么？", "用户登录后服务端签发包含用户标识和过期时间的 JWT；客户端携带 Bearer Token；服务端验证签名、过期时间和权限后允许访问。JWT 不应存放敏感明文。", "当前 code_helper 尚未接入用户认证，接口默认运行在本地；如果对外部署，应在 Depends 中加入 Token 验证和用户权限。", "继续说明 access token、refresh token、密钥轮换和撤销策略。"),
        q("OAuth2 是什么？", "OAuth2 是授权框架，描述第三方或客户端如何获得访问资源的令牌。它不等于登录协议；OpenID Connect 在 OAuth2 上补充身份认证信息。", "当前项目是本地开发工具，不需要复杂 OAuth2；多人部署时可接入公司 SSO 或 OIDC。", "不要把 OAuth2、JWT、OIDC 混为同一个概念。"),
        q("密码为什么不能直接保存？", "密码应使用带盐的慢哈希算法，如 Argon2 或 bcrypt。服务端只保存哈希，登录时对输入密码做同样哈希并比较，不能解密还原。", "当前项目没有用户密码表；如果增加账号系统，应把认证逻辑与 TaskService 分离，并对 API Key、密码和 Token 脱敏。", "MD5、SHA1 不是密码哈希方案，盐值和工作因子同样重要。"),
        q("CORS 是什么？", "CORS 是浏览器的跨源访问控制。后端通过响应头声明哪些前端源可以访问，防止任意网页读取接口响应。它不是身份认证，也不能替代服务端权限校验。", "main.py 使用 CORSMiddleware，默认允许 localhost:5173 的前端访问后端 8000。", "生产环境不能长期使用 allow_origins=['*'] 配合凭据，应明确配置可信来源。"),
        q("如何防止 SQL 注入？", "使用参数化查询或 ORM，不要把用户输入拼接进 SQL 字符串；同时限制数据库账号权限并校验输入。", "SQLiteStore 的查询使用问号参数，例如按 task_id 查询；搜索记忆的关键词在 Python 中评分，没有直接拼接 SQL。", "参数化查询解决注入，不能替代权限控制、日志脱敏和输入长度限制。"),
        q("如何防止路径穿越？", "对用户提供的路径做规范化，解析为绝对路径，然后确认它位于允许的工作区根目录内；拒绝 ..、符号链接逃逸和工作区外写入。", "CodeTaskWorkflow._resolve_workspace_path 与 assistant_tools 的路径检查确保补丁只能作用于用户选择的 workspace。", "需要讨论符号链接、大小写路径、Windows 盘符和 TOCTOU 风险。"),
        q("文件上传有哪些风险？", "要限制文件大小和类型、生成安全文件名、避免直接执行、扫描恶意内容、校验存储路径，并限制下载权限。", "当前项目不是通用文件上传服务，主要由 Agent 在受控 workspace 内读取和写入文本文件；写入工具仍会校验路径和已有文件。", "如果以后支持上传，应将上传目录与可执行工作区隔离。"),
        q("为什么测试命令要做白名单？", "Agent 生成的命令不能直接交给 Shell，否则可能删除文件、访问系统或泄露密钥。白名单应同时限制可执行文件和参数前缀。", "settings.py 默认只允许 python -m pytest -q；assistant_tools.run_command 会校验命令，测试还受 workspace_root 和超时限制。", "这是 code_helper 的核心安全设计，回答时要明确“LLM 可以建议，系统决定是否执行”。"),
        q("如何保护 API Key 和敏感信息？", "密钥放环境变量或密钥管理服务，不提交到 Git，不返回给前端；日志、错误和 LLM 上下文中也要脱敏。", "LLMConfig 使用 api_key_env，从环境读取 OPENAI_API_KEY；前端只收到回复、状态和必要结果。", "还要考虑 .env 权限、日志过滤、错误堆栈和第三方 LLM 数据保留策略。"),
        q("Agent 如何保证不能修改工作区之外的文件？", "LLM 只生成候选路径，系统必须重新解析路径、校验工作区边界、校验旧文本、限制创建和修改语义，不能相信 LLM 声明。", "Executor 的真正写入由 assistant_tools.write_text_file 和 workflow 的路径解析完成；补丁前创建快照，Reviewer 再检查 changed_files 和测试。", "安全边界必须位于工具层，而不是只写在 Prompt 中。"),
    ]),
    ("七、测试、部署与实时通信", [
        q("如何测试 FastAPI 接口？", "可以用 pytest 加 TestClient 或 httpx AsyncClient，覆盖状态码、响应模型、异常路径和数据库变化。外部服务应使用 mock，数据库可以使用临时库。", "apps/api/tests/test_api.py 测试创建任务、后台运行、最终状态、事件、记忆、回滚、聊天追问、命令隔离和 SSE。", "测试异步后台任务时要轮询终态并设置超时，不能无限 sleep。"),
        q("如何替换测试数据库？", "通过配置注入数据库路径或 Store，在 fixture 中创建临时 SQLite 数据库；测试结束清理文件。不要让测试依赖开发环境的真实 app.db。", "当前测试 reset_state 会清空 data/app.db、snapshots 和 demo workspace；更完善的版本可以把数据库路径通过 Settings 依赖替换为 tmp_path。", "测试需要隔离数据库、文件系统、快照和 LLM 响应。"),
        q("如何测试异步接口？", "使用 pytest-asyncio 或 AnyIO，再用 httpx.AsyncClient；需要测试后台任务时等待可观察的状态变化，并验证异常和取消。", "run_task 返回 202 后，测试通过 wait_for_terminal 轮询 GET /tasks/{id}，确认后台 Agent 最终成功或失败。", "避免用固定长时间 sleep，优先轮询条件并设置 deadline。"),
        q("为什么开始执行接口返回 202？", "202 Accepted 表示请求已接受，但处理尚未完成。它适合长时间任务，避免 HTTP 请求一直占用连接；客户端随后通过查询或事件流获取最终状态。", "POST /tasks/{task_id}/run 启动 asyncio 后台任务，返回 accepted；前端使用 EventSource 和 5 秒轮询跟踪进度。", "如果请求需要立即完成并返回结果，应使用 200；不要用 202 掩盖同步错误。"),
        q("为什么使用 SSE 而不是 WebSocket？", "SSE 是服务器到浏览器的单向文本事件流，浏览器原生支持 EventSource，适合任务进度和日志推送；WebSocket 是双向通信，适合实时协作、双向控制和高频交互。", "code_helper 的用户输入走普通 POST，后端进度走 GET /tasks/{task_id}/events 的 SSE，职责清晰。", "SSE 需要处理断线、重连、Last-Event-ID 或 sequence；当前项目用 sequence 加 5 秒轮询兜底。"),
        q("如何选择 Uvicorn Worker 数量？", "Worker 是独立进程，通常根据 CPU、I/O、内存和压测结果确定，而不是固定套公式。每个 Worker 都会有自己的内存、连接池和进程内任务。", "当前 TaskService 的 asyncio 任务和订阅队列是进程内状态，适合单 Worker；多 Worker 部署需要外部队列和事件总线。", "要结合数据库连接数、LLM 并发、SSE 连接数和容器资源回答。"),
        q("为什么生产环境需要反向代理？", "Nginx、Caddy 或云负载均衡可以处理 TLS、域名、静态文件、压缩、限流、超时和转发，把应用服务器与公网隔离。", "开发环境前端 Vite 5173、后端 Uvicorn 8000 分开运行；生产环境可以由反向代理统一域名并转发 /api 和 SSE。", "SSE 代理必须关闭不当缓冲并设置较长读取超时。"),
        q("Docker 部署要注意什么？", "镜像应固定 Python 版本和依赖，使用非 root 用户，配置环境变量，分离代码、数据库和工作区卷，并设置健康检查。", "code_helper 的 data/app.db、snapshots 和用户 workspace 需要持久化卷，否则容器重建会丢失任务和文件。", "还要考虑镜像大小、启动命令、日志输出和文件权限。"),
        q("如何做健康检查？", "健康检查至少区分进程存活和依赖可用：liveness 判断进程是否能响应，readiness 判断数据库、队列等是否准备好。", "main.py 提供 /health 返回 status=ok；生产版本可以额外检查 SQLite、LLM 配置和任务队列，但不应把慢检查放进每次探活。", "健康检查响应不应泄露密钥、路径或内部堆栈。"),
        q("如何做日志、监控和错误追踪？", "使用结构化日志记录 request_id、task_id、耗时、状态和错误；指标记录请求延迟、错误率、任务成功率、LLM 延迟和队列长度；异常通过 Sentry 或 OpenTelemetry 追踪。", "code_helper 已有 task_id、event sequence、task status 和 test duration 等可观测字段，后续可以统一成结构化日志和 trace_id。", "不要把完整 prompt、API Key 或用户源码无控制地写入日志。"),
        q("如果服务重启，后台任务能恢复吗？", "当前进程内 asyncio.create_task 会丢失执行过程，但任务和事件已经持久化。生产系统需要持久化任务队列、幂等执行、租约/锁和恢复扫描。", "code_helper 适合本地 MVP；服务重启后可以看到数据库中最后状态，但不能保证正在执行的 Agent 自动续跑。", "这是从 MVP 到生产系统的重要差距。"),
        q("如果有十万个任务，当前设计要如何改？", "需要分页查询、数据库索引、独立队列、水平扩展、事件存储分层、Redis 或消息队列、批量写入和归档策略。SSE 连接也要通过共享事件总线支持多实例。", "当前 SQLite 和进程内订阅集合适合小规模本地运行；迁移 PostgreSQL、Redis、Celery、对象存储和统一事件总线是自然演进路径。", "先明确瓶颈，再用压测数据决定优化，不能只堆框架。"),
    ]),
    ("八、code_helper 项目架构与 Agent 追问", [
        q("请用 2 分钟介绍 code_helper。", "这是一个面向代码工作区的 R&D Assistant。前端 React 通过 HTTP JSON 创建任务、发送聊天消息，通过 SSE 接收进度；后端 FastAPI 负责路由和业务服务，Agent 检索文件、规划、调用受控工具写入代码、运行白名单测试并审查结果；SQLite 保存任务、事件、记忆和产物，快照用于回滚。", "主链路是 App.tsx → api/client.ts → FastAPI routes → TaskService → CodeTaskWorkflow/LangGraph → assistant_tools/SQLiteStore。", "回答最后补充限制：当前适合单机 MVP，多 Worker、持久化队列和多用户权限仍需加强。"),
        q("用户点击 Run task 后完整发生什么？", "前端 POST /tasks/{id}/run；后端设置 queued 并创建后台任务；TaskService 准备上下文和记忆，LangGraph 执行 retrieval、planner、executor、reviewer、rollback/finish；每个关键节点写事件；前端通过 SSE 和轮询更新；最终保存 diff、测试、记忆和终态。", "TaskService._execute 通过 run_task_lifecycle_graph 组织 prepare、execute、persist；CodeTaskWorkflow.run 负责任务内部的检索、子目标和 Agent 图。", "重点说明接口返回 202，不代表已经完成。"),
        q("为什么要有 Service 和 Storage 分层？", "路由层负责 HTTP 协议，Service 层负责业务流程，Storage 层负责持久化。这样可以独立测试业务、替换数据库、避免路由过于复杂。", "tasks.py 很薄，只负责调用 TaskService 和把异常转 HTTP；TaskService 处理任务、事件、记忆和后台执行；SQLiteStore 处理 SQL。", "如果把所有代码写在路由函数中，测试、复用和未来迁移数据库都会变困难。"),
        q("为什么使用 LangGraph 编排 Agent？", "LangGraph 用显式 StateGraph 表达节点、状态和条件边，适合需要循环、回滚、人工审批、暂停恢复和流式事件的 Agent 流程。它比把所有逻辑塞进一个 while 循环更容易观察和扩展。", "workflow/langgraph_workflow.py 定义 AgentGraphState，包含 retrieval_context、planner、approval_gate、executor、reviewer、rollback、finish；聊天也有 chat_agent、route 和 chat_reviewer 子图。", "要强调 LangGraph 只编排，文件安全、测试白名单和数据库仍由确定性代码控制。"),
        q("Planner、Executor、Reviewer 分别做什么？", "Planner 决定下一步和候选补丁；Executor 选择受控工具并执行文件读写、测试；Reviewer 根据真实补丁、测试结果和用户要求判断 approve、revise 或 reject。", "Planner 通过 choose_parallel_branch 和 OpenAIPlanner 生成候选；Executor 使用 plan_executor_tool；Reviewer 使用 review_result；拒绝后进入 rollback。", "LLM 的声明不能替代实际文件、测试和 changed_files 证据。"),
        q("Agent 之间如何共享状态？", "LangGraph StateGraph 使用一个类型化状态对象在节点之间传递字段。节点返回增量更新，条件路由根据 next_node 或状态字段决定下一节点。", "AgentGraphState 保存 task_id、retrieval_hits、current_context、decision、proposal、final_test、review、branch_history、snapshot_path 和 summary。", "状态应避免塞入过大的源码或敏感信息；可用 checkpoint、数据库或对象存储保存大对象。"),
        q("Reviewer 要求修改时如何回滚？", "Reviewer 返回 revise/reject，图路由到 rollback。rollback 恢复快照、清空当前补丁和 changed_files、读取额外文件并把反馈传回 Planner。", "_rollback_node 调用 restore_snapshot，并发出 branch.reverted，再回到 planner；这样错误补丁不会污染下一轮。", "回滚还需要验证快照完整性，防止快照被篡改。"),
        q("为什么 Agent 不能直接执行任意命令？", "LLM 输出不可信，任意命令可能删除文件、读取密钥或访问网络。正确做法是工具白名单、参数校验、工作区边界、超时和审计事件。", "assistant_tools.run_command 只允许配置中的 pytest 命令；Executor 的工具选择必须落到系统实现，不能直接执行 LLM 返回的 Shell 字符串。", "把安全约束放在工具实现里，而不是只写在 Prompt 里。"),
        q("如何保证最终回复真正回答用户要求？", "最终摘要不能只复述 Agent 状态，应读取 task_description、实际 changed_files、测试结果和审查结果，覆盖用户所有明确交付项。Chat Agent 还要经过 chat_reviewer 检查，防止把计划说成完成。", "LLM Prompt 要求 summary 覆盖复杂度等显式要求；TaskService._build_terminal_summary 汇总最终结果，前端在终态用 task.summary 替换过程中的占位回复。", "可继续追问如何处理没有文件变化、测试未收集和 Reviewer 未批准等边界。"),
        q("如何支持人工审批？", "在产生可写补丁后插入 interrupt/approval gate，向前端发出审批事件；用户批准则继续 Executor，拒绝则 Finish 或记录 reject。需要 checkpointer 才能可靠恢复。", "LangGraph 图中的 approval_gate 使用 interrupt，AgentGraphState 有 require_human_approval、approval_response 字段；默认是否开启由任务配置决定。", "审批不能绕过路径校验和测试，人工批准也只批准候选补丁。"),
        q("LangGraph 不可用时怎么办？", "框架依赖不可用时应让应用仍能启动，使用兼容的原生流程；同时记录 fallback 事件或日志，避免用户误以为使用了图编排。", "langgraph_workflow.py 定义 LangGraphUnavailable，run_subgoal_graph、run_task_lifecycle_graph 和聊天图都提供原生回退路径。", "回退路径必须通过相同的安全工具和状态持久化，不能变成一套更宽松的执行逻辑。"),
    ]),
    ("九、现场编码题与标准思路", [
        q("现场写一个 CRUD 接口怎么做？", "先定义 Pydantic 输入和输出模型，再定义路由和状态码，Service 负责业务，Repository 负责数据库。处理 404、校验错误和事务，最后补测试。", "可仿照 tasks.py：POST /tasks 创建任务，GET /tasks 列表和详情，DELETE /tasks/{id} 删除；路由不直接写 SQL。", "先写最小正确版本，再补鉴权、分页、日志和测试，不要一开始引入过多抽象。"),
        q("如何实现分页查询？", "使用 page/page_size 或 limit/offset，并限制最大 page_size；查询总数和当前页数据，按稳定字段排序，必要时使用游标分页。", "任务列表当前是简单 list_tasks；任务量变大时应在 SQLiteStore 增加 LIMIT/OFFSET 或基于 created_at/id 的游标，并把分页元数据放入响应模型。", "大数据量场景优先考虑游标分页，避免深 offset 扫描。"),
        q("如何写 JWT 登录接口？", "定义 LoginRequest，校验用户密码，生成带 sub、exp、iat 的签名 Token；用 OAuth2PasswordBearer 读取 Token，get_current_user 依赖验证并注入路由。", "当前项目没有用户表，但可以在 SQLite 增加 users 表，并把认证依赖接入 tasks 路由；任务还要增加 owner_id 做权限过滤。", "不要把密码或完整 Token 写入日志和事件 payload。"),
        q("如何写文件上传接口？", "使用 UploadFile，限制 Content-Length、文件名、扩展名和大小，将内容写入隔离目录并进行安全命名；不要直接执行上传文件。", "code_helper 的 workspace 是 Agent 的受控工作区，不能直接把它当作上传目录；上传区和可执行源码区应隔离。", "还要考虑病毒扫描、断点上传、下载权限和路径穿越。"),
        q("如何实现统一异常处理？", "定义业务异常类型和状态码映射，在 app.exception_handler 中统一返回 error_code、message、request_id；未知异常只返回通用错误并记录服务端日志。", "当前 tasks.py 对 KeyError、RuntimeError、ValueError 做了局部转换；进一步可以统一封装 NotFound、Conflict、WorkspaceSecurityError 和 WorkflowError。", "响应结构稳定比把 Python 异常字符串直接返回更重要。"),
        q("如何保证一次任务的数据库操作一致？", "为单个状态转换和事件写入定义事务边界，异常时回滚；文件操作和外部调用不要长时间占用数据库事务，使用状态机和补偿逻辑处理跨系统一致性。", "任务状态写入 SQLite，文件修改通过 snapshot/restore 补偿；LangGraph 节点完成后发事件，最终由 TaskService 持久化结果。", "这属于分布式一致性问题，不能用一个数据库事务包住文件系统和 LLM。"),
        q("如何把耗时任务放到后台？", "短任务可以使用 BackgroundTasks；需要可靠重试、跨进程和持久化的任务应使用 Celery、RQ 或云队列。API 返回 202，并提供状态查询和事件流。", "TaskService 使用 asyncio.create_task 作为单机 MVP 实现；未来可迁移为 Redis/Celery，保留同样的任务状态接口。", "必须处理幂等、取消、重试次数、死信和任务超时。"),
        q("如何实现 SSE 接口？", "返回 text/event-stream 的 StreamingResponse，服务端按 event/data/id 格式发送并及时 flush；客户端使用 EventSource，处理断开和重连。", "events.py 的 generator 从 EventService.stream 获取事件，前端 TaskDetail 使用 EventSource 连接 /tasks/{id}/events。", "需要 sequence 或 Last-Event-ID 防止断线后重复或漏事件。"),
        q("如何用 Depends 实现权限控制？", "先通过认证依赖获取用户，再通过权限依赖检查角色或资源所有权，最后注入到路由。认证和授权要分开。", "未来可在 TaskService.get_task 前校验 task.owner_id；当前 get_container 只负责获取应用容器，没有用户权限。", "不能只在前端隐藏按钮，后端每次请求都必须校验。"),
        q("如何写一个 FastAPI 接口测试？", "准备 fixture 和测试数据，调用 TestClient/AsyncClient，断言状态码、响应字段和副作用。对后台任务轮询终态，对外部 LLM 使用 mock。", "test_api.py 会创建任务、调用 /run、用 wait_for_terminal 查询终态，并检查 diff、测试、memory、events、rollback 和 SSE。", "测试不仅断言 200，还要验证错误状态、事件顺序和文件是否真的改变。"),
    ]),
]


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "FastAPI 面试问答 · code_helper 实战版"
    for run in header.runs:
        set_font(run, 8.5, GRAY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(95)
    cover.paragraph_format.space_after = Pt(12)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("FASTAPI INTERVIEW GUIDE")
    set_font(run, 11, BLUE, True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("FastAPI 面试问答")
    set_font(run, 30, NAVY, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("结合 code_helper 项目的原理、实践与追问答案")
    set_font(run, 14, GRAY)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(90)
    run = meta.add_run("面向：FastAPI 后端开发岗位\n内容：原理问答 · 项目表达 · 现场编码 · 生产化思考")
    set_font(run, 10.5, GRAY)
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    shade_cell(callout.cell(0, 0), LIGHT_BLUE)
    p = callout.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("使用建议：先读项目主链路，再按问题练习口头回答，最后用“当前设计 → 原因 → 限制 → 改进”结构应对追问。")
    set_font(run, 10.5, NAVY, True)
    doc.add_page_break()

    doc.add_heading("使用说明", level=1)
    add_body(doc, "这份文档不是 FastAPI API 手册，而是一份面试表达手册。每道题都包含标准回答、code_helper 中的对应实现和面试追问方向。回答时先给结论，再结合项目证据，最后主动说明当前限制和生产环境改进。")
    doc.add_heading("项目速览", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    for cell, text in zip(table.rows[0].cells, ("组件", "code_helper 中的对应内容")):
        shade_cell(cell, LIGHT_BLUE)
        cell.paragraphs[0].add_run(text)
        for run in cell.paragraphs[0].runs:
            set_font(run, 10, NAVY, True)
    overview = [
        ("前端", "React + TypeScript + Vite；fetch 调用 JSON API，EventSource 接收 SSE，5 秒轮询兜底。"),
        ("后端", "FastAPI + Uvicorn；main.py 注册路由、中间件和应用容器。"),
        ("业务层", "TaskService 管理任务、聊天、后台执行、记忆和最终持久化。"),
        ("Agent 编排", "LangGraph 状态图驱动检索、Planner、Executor、Reviewer、Rollback 和 Finish；聊天有独立子图。"),
        ("存储", "SQLite 保存任务、事件、子目标、记忆、快照和产物；默认 data/app.db。"),
        ("安全", "工作区路径边界、补丁 old 文本匹配、快照完整性、命令白名单和测试超时。"),
    ]
    for label, detail in overview:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for index, cell in enumerate(cells):
            for run in cell.paragraphs[0].runs:
                set_font(run, 10, DARK_BLUE if index == 0 else RGBColor(0, 0, 0), index == 0)
    set_table_geometry(table, [2700, 6660])
    doc.add_heading("项目关键文件", level=1)
    for path, purpose in [
        ("apps/api/main.py", "创建 FastAPI 应用、注册 CORS、初始化 SQLite 和 Service。"),
        ("apps/api/api/routes/tasks.py", "任务、聊天、运行、回滚、删除和产物 API。"),
        ("apps/api/services/task_service.py", "任务后台执行、LangGraph 生命周期、聊天图、记忆和结果持久化。"),
        ("packages/agent-core/src/agent_core/workflow/langgraph_workflow.py", "AgentGraphState、任务图、聊天图、子目标图、回退和人工审批节点。"),
        ("packages/agent-core/src/agent_core/llm.py", "Planner、Executor tool router、Reviewer 和 Chat Agent 的 LLM 调用。"),
        ("packages/tools/src/assistant_tools/", "文件读写、快照、回滚、命令白名单和 pytest 工具。"),
        ("apps/web/src/api/client.ts", "前端 HTTP API 和 SSE URL 封装。"),
        ("apps/web/src/components/TaskDetail.tsx", "任务详情、对话、过程卡片、SSE 和轮询刷新。"),
    ]:
        add_body(doc, f"{path}：{purpose}")
    doc.add_page_break()

    total = 0
    for section_title, questions in SECTIONS:
        doc.add_heading(section_title, level=1)
        for question, answer, project, followup in questions:
            total += 1
            add_q(doc, total, question, answer, project, followup)

    doc.add_heading("附录：面试前检查清单", level=1)
    for text in [
        "能用两分钟讲清 React → FastAPI → TaskService → LangGraph → SQLite/SSE 的完整链路。",
        "能解释 async、事件循环、阻塞代码、to_thread 和多 Worker 的关系。",
        "能写出带 Pydantic、Depends、异常处理和测试的 CRUD 接口。",
        "能解释 202、SSE、断线重连、事件 sequence 和 5 秒轮询兜底。",
        "能说明文件路径校验、快照回滚和命令白名单为什么必须由确定性代码执行。",
        "能主动说出当前 MVP 的限制：SQLite、进程内后台任务、单机事件订阅和认证缺失。",
        "能给出生产化方向：PostgreSQL、Redis/Celery、OIDC、反向代理、可观测性和多实例事件总线。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text)
        set_font(run, 10.5)
    add_body(doc, f"本手册共整理 {total} 道问题，答案以当前工作区代码为依据；代码持续演进时，应同步更新项目速览和项目追问部分。", bold_label="版本说明：")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
