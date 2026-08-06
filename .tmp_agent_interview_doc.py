from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "docs" / "Agent面试问答_CodeHelper实战版.docx"
FONT = "Microsoft YaHei"
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"


def set_font(run, size=10.5, color=None, bold=None, italic=None):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("eastAsia", "ascii", "hAnsi"):
        fonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    set_font(tail, 9, GRAY)


def body(doc, text, label=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    if label:
        r = p.add_run(label)
        set_font(r, 10.5, DARK_BLUE, True)
    r = p.add_run(text)
    set_font(r, 10.5, color or RGBColor(0, 0, 0))
    return p


def question(doc, number, item):
    title, answer, project, followup = item
    h = doc.add_heading(f"Q{number}. {title}", level=2)
    h.paragraph_format.keep_with_next = True
    body(doc, answer, "标准回答：")
    body(doc, project, "结合 code_helper：")
    body(doc, followup, "面试追问：", GRAY)


def q(title, answer, project, followup):
    return title, answer, project, followup


SECTIONS = [
    ("一、LLM Agent 基础概念", [
        q("什么是 LLM Agent？", "LLM Agent 是以大语言模型作为决策组件、能够观察环境、规划下一步并调用工具执行动作的系统。它通常由模型、状态、工具、记忆、编排和安全边界组成，不等于一次普通的聊天补全。", "code_helper 的 Planner 读取任务与代码上下文，选择 read_more、patch 或 finish；Executor 调用受控工具；Reviewer 根据真实结果决定 approve、revise 或 reject。", "继续区分聊天机器人、工作流和 Agent：关键在于是否根据环境反馈自主选择动作并形成闭环。"),
        q("Agent 和普通 Chatbot 有什么区别？", "Chatbot 主要生成文本回复；Agent 还会维护任务状态、调用外部工具、接收执行结果并再次决策。Agent 的价值是完成任务，而不仅是回答问题。", "Chat Agent 负责对话和意图判断，代码任务 Agent 则能读取文件、写补丁、运行 pytest、审查和回滚。", "面试时强调 Agent 的自主性必须受到工具权限、步数和预算约束。"),
        q("Agent 和传统工作流有什么区别？", "传统工作流路径通常由开发者预先确定；Agent 在部分节点由 LLM 根据上下文选择动作。可靠系统往往是两者结合：确定性框架控制边界，LLM 只负责需要语义判断的部分。", "LangGraph 固定 retrieval、planner、executor、reviewer、rollback、finish 节点，Planner 和 Reviewer 的具体判断由 LLM 产生。", "不要把所有逻辑都交给模型；权限、文件写入、状态迁移和持久化应保持确定性。"),
        q("Agent 的核心组件有哪些？", "常见组件包括模型、系统提示词、工具、状态、短期上下文、长期记忆、规划器、执行器、评审器、事件与评测系统。复杂系统还需要检查点、人工审批和预算控制。", "项目中 OpenAIPlanner 是模型接口，assistant_tools 是工具层，AgentGraphState 是状态，SQLite task_memory 是长期记忆，LangGraph 是编排层。", "能否清楚画出组件边界，通常比背某个框架 API 更重要。"),
        q("什么是 Agent Loop？", "Agent Loop 是观察、思考/决策、行动、获取结果、再决策的循环。循环必须有终止条件，例如完成、失败、达到最大轮数、预算耗尽或人工终止。", "Planner 生成动作，Executor 执行，Reviewer 返回反馈；revise/reject 经 rollback 回到 Planner，max_turns 默认限制循环次数。", "说明无限循环、重复工具调用和成本失控如何防止。"),
        q("什么是 ReAct？", "ReAct 将 reasoning 与 acting 结合：模型基于观察选择动作，系统执行工具并把 observation 返回，模型再决定下一步。产品中不必展示隐藏思维链，只需要结构化动作和可观察进度。", "ReActDecision 只暴露 action、summary、rationale、files_to_read 和 proposal，不把隐藏思考链显示给前端。", "可以追问为什么结构化 rationale 不等同于完整 chain-of-thought。"),
        q("Agent 什么时候不值得使用？", "任务路径固定、规则清晰、容错要求极高或成本延迟敏感时，普通程序和状态机更合适。Agent 应用于语言理解、模糊规划、非结构化检索等确定性规则难覆盖的部分。", "测试命令白名单、路径校验、SQLite 写入和快照恢复都没有交给 LLM；LLM 主要用于规划、工具路由、审查和聊天。", "回答重点是判断力：不是所有系统都应该“Agent 化”。"),
        q("单 Agent 和多 Agent 如何选择？", "单 Agent 更简单、延迟低、上下文一致；多 Agent 适合角色冲突、并行探索、独立审查或不同工具权限。引入多 Agent 前要证明它提升质量，而不是只增加调用次数。", "code_helper 使用 Planner、Executor、Reviewer、Chat Agent，并在 Planner 内并行生成 balanced、critic、memory、explorer、verifier 候选。", "多 Agent 最大问题通常是成本、状态一致性、重复劳动和责任边界。"),
    ]),
    ("二、Prompt、结构化输出与上下文", [
        q("系统提示词应该包含什么？", "系统提示词应明确角色、目标、允许动作、输入证据、输出 schema、安全限制、失败策略和终止条件。约束应可验证，不能只写抽象口号。", "Planner Prompt 要求 JSON、限定 read_more/patch/finish；Executor 限定四种工具；Reviewer 限定 approve/revise/reject，并要求依据 changed_files 和测试结果。", "真正的安全约束仍必须在工具层实现，Prompt 只能减少错误，不能构成权限边界。"),
        q("为什么要使用结构化输出？", "结构化输出让程序可以稳定解析动作、路径、补丁和审查决定，减少从自然语言中猜测意图。可以使用 JSON Schema、Pydantic 或模型原生 structured outputs。", "llm.py 把模型输出解析为 ReActDecision、ExecutorToolCall、ReviewDecision 和 ChatResponseResult。", "还要处理字段缺失、非法枚举、JSON 外多余文本和模型拒答。"),
        q("如何处理模型返回非法 JSON？", "先使用严格 schema 和低温度；解析失败时记录原始错误、有限重试或回退到安全策略。不能因为解析失败就猜测并执行高风险动作。", "OpenAIPlanner 解析失败会返回 None，工作流转入 heuristic/fallback 或要求重新规划，不直接写文件。", "重试要限制次数并避免把敏感原始输出写入普通日志。"),
        q("上下文窗口如何管理？", "只保留任务相关信息，对代码做检索和截断，对历史做摘要，对工具结果做结构化压缩。上下文越大并不一定越好，噪声会降低决策质量并增加成本。", "项目先 search_workspace_files，再只读取 focus_paths、retrieval hits、测试文件和相关 memory；传给 LLM 的文件内容有长度限制。", "继续讨论 token 预算、优先级、去重和上下文污染。"),
        q("如何避免 Prompt Injection？", "把用户文本、仓库文件和工具输出都视为不可信数据；系统指令与数据分离；工具权限由代码限制；对敏感操作增加审批；不要允许检索内容覆盖系统规则。", "工作区源码可能包含诱导 Agent 执行命令的文字，但 Executor 只能选择受控工具，run_command 还会执行白名单校验。", "Prompt Injection 不能只靠“忽略恶意指令”一句提示词解决。"),
        q("如何让最终回答真正覆盖用户要求？", "在任务状态中保存原始需求，Reviewer 检查所有明确交付项，最终摘要基于真实变更和测试证据生成。对复杂度、文件名等要求要逐项校验。", "Reviewer Prompt 明确要求 summary 回答 task_description 的所有交付项；TaskService 再生成 terminal summary，并让 Chat Reviewer 复核对话回复。", "可追问如何用 checklist 或 requirement extraction 做自动覆盖率评测。"),
        q("温度、最大输出和超时如何设置？", "规划和工具调用需要稳定结构，通常使用较低温度；max tokens 要足以容纳完整文件或补丁；超时要结合模型延迟和任务 SLA，并设置重试与熔断。", "LLMConfig 配置 temperature、max_output_tokens、timeout_seconds，模型与密钥从配置和环境读取。", "参数应通过评测和线上指标调整，而不是照搬固定值。"),
        q("如何设计多轮对话上下文？", "保留用户消息、已确认执行结果、当前任务状态和最近关键事件；区分对话意图与代码执行意图；不要把旧任务的未验证计划当作事实。", "TaskService._build_chat_context 汇总 task、events、artifacts 和结果；Chat Agent 分类 implementation、question、panel、status 等意图。", "多轮上下文需要避免重复执行、消息顺序错乱和历史无限增长。"),
    ]),
    ("三、工具调用与安全执行", [
        q("什么是 Tool Calling？", "Tool Calling 是模型输出结构化工具名称和参数，由宿主程序校验并执行，再把结果返回模型。模型没有直接系统权限，真正权限属于工具实现。", "ExecutorToolCall 支持 read_file、write_file、run_tests、finish，执行逻辑位于 LangGraph executor 节点和 assistant_tools。", "要强调模型选择工具不等于工具一定被允许执行。"),
        q("如何设计一个好的 Agent 工具？", "工具应职责单一、参数明确、结果结构化、幂等性可说明、错误可观察，并限制作用域、超时和资源消耗。描述要让模型知道何时使用和何时不能使用。", "read/write/test 分开，文件路径必须 workspace-relative，write_file 区分 create/modify，run_tests 只执行允许命令。", "避免设计万能 shell 工具或参数含义模糊的大工具。"),
        q("为什么不能给 Agent 任意 Shell？", "LLM 输出可能错误或被注入，任意 Shell 会扩大到删除文件、读取密钥、联网和持久化攻击。最小权限原则要求只暴露完成任务所需的工具。", "code_helper 的命令层只允许 python -m pytest -q，测试受 workspace_root、参数前缀和超时约束。", "即使是本地开发工具，也要假设仓库内容和用户输入不可信。"),
        q("如何校验文件写入？", "规范化路径并确认位于工作区；创建时禁止覆盖；修改时要求 old 文本准确匹配；限制文件类型和大小；写入前快照，写入后生成 diff。", "executor 节点调用 _resolve_workspace_path、replace_once、write_text_file 和 artifact_from_patch，changed_files 来自实际目标路径。", "还需考虑符号链接、路径大小写和竞态条件。"),
        q("工具执行失败后怎么办？", "工具应返回明确错误，不伪造成功；Reviewer 根据错误决定 revise/reject；必要时恢复快照并把反馈送回 Planner。重试必须有限且避免重复副作用。", "executor 捕获异常写入 execution_error，Reviewer 生成 reject，rollback 恢复 workspace 后回到 Planner。", "区分可重试错误、参数错误、权限错误和不可恢复错误。"),
        q("如何处理工具幂等性？", "读工具天然幂等；写工具需要 operation、版本或预期旧值；任务系统使用 idempotency key、状态检查和事务避免重复副作用。", "modify 要求 old 精确匹配，因此同一补丁重复执行会失败而不是重复修改；活动任务还通过 _running_tasks 避免重复启动。", "生产多 Worker 场景需要分布式锁或数据库唯一约束。"),
        q("为什么写入前要创建快照？", "LLM 补丁具有不确定性，快照提供文件系统级补偿机制，使 Reviewer 拒绝、测试失败或人工回滚时可以恢复。", "executor 在 branch/before_patch 创建 snapshot，rollback 节点调用 restore_snapshot，快照 manifest 还会校验完整性。", "数据库事务不能自动回滚文件系统，因此需要额外补偿。"),
        q("人工审批应该放在哪里？", "高风险动作应在计划生成后、真正执行前 interrupt；审批内容包含路径、操作、diff 和风险；批准后仍需正常安全校验。", "Agent 图有 approval_gate 和 require_human_approval，使用 LangGraph interrupt 支持批准或拒绝候选补丁。", "要有 checkpointer 才能可靠暂停和恢复，审批不能只存在进程内。"),
    ]),
    ("四、多 Agent 角色与协作", [
        q("为什么拆成 Planner、Executor、Reviewer？", "三个角色分别负责语义规划、受控执行和独立验收，降低单次模型调用既提出方案又自我批准的偏差。角色拆分还便于分配不同权限、模型和评测指标。", "Planner 不能直接写文件；Executor 只能调用工具；Reviewer 依据实际 patch、changed_files 和 TestOutcome 审查。", "角色拆分不自动带来质量，必须定义清晰输入、输出和责任边界。"),
        q("Planner 应输出什么？", "Planner 应输出下一动作、必要文件、完整候选补丁、用户可见摘要和简短理由。它不应输出无法执行的伪代码或隐藏思维链。", "ReActDecision 包含 action、files_to_read、proposal、summary、rationale、provider 和 model。", "高质量 Planner 的关键指标是补丁可执行率、任务覆盖率和无效读取次数。"),
        q("Executor 是否一定需要 LLM？", "确定性工具执行本身不需要 LLM；LLM 可以用于选择下一工具或补全参数，但最终执行必须由代码校验。简单任务可以完全由状态机执行。", "code_helper 的 Executor Tool Router 由 LLM 选择 read/write/test/finish，但文件读写和 pytest 都是确定性函数。", "回答时避免把“模型调用工具”和“模型直接执行工具”混为一谈。"),
        q("Reviewer 如何保持独立？", "Reviewer 应获取原始需求、实际 diff、测试、变更文件和执行错误，而不是只读取 Planner 摘要。可以使用不同 Prompt、不同模型或规则检查增强独立性。", "review_result 的状态包含 task_description、patch_applied、changed_files、latest_test 和 branch_history；测试失败不能 approve。", "同一个模型仍可能产生相关偏差，可加入静态检查、测试和人工审批。"),
        q("Chat Agent 的职责是什么？", "Chat Agent 负责意图分类、解释状态、回答问题和触发后续实现任务；它不能自行声称文件已修改，必须依据事件与任务结果。", "run_chat_graph 包含 chat_agent、implementation/panel/informational route 和 chat_reviewer；实现请求会排队启动 _execute。", "对话 Agent 与代码执行 Agent 应分离权限。"),
        q("为什么要有 Critic、Explorer、Verifier 等候选？", "不同 profile 引导模型从保守、广泛检索、记忆和测试角度提出候选，减少单一路径的盲点。并行候选适合高不确定任务，但会增加成本。", "parallel_branches.py 使用 Send 并行构建 planner、critic、memory、explorer、verifier 五个分支，再根据 score 选择。", "应通过评测证明多分支比单 Planner 更好，并设置候选数量上限。"),
        q("多 Agent 如何避免互相重复？", "给每个角色明确目标和只读/可写权限，状态中记录已读文件、已拒绝补丁和分支历史，路由器避免重复动作，并使用统一任务 ID。", "review_feedback 和 branch_history 会回传 Planner，Prompt 要求不要重复已拒绝补丁。", "还可加入内容哈希、动作去重和成本预算。"),
        q("多 Agent 如何解决冲突？", "冲突不能靠模型无限辩论，应由可验证证据、测试、规则评分或最终 Reviewer 决定。高风险冲突交给人工审批。", "候选分支根据是否有完整 patch、测试状态、读取测试文件和 profile 加权；Reviewer 对选中分支做最终审查。", "评分函数也可能偏置，应记录候选和分数用于离线分析。"),
    ]),
    ("五、LangGraph 编排", [
        q("为什么选择 LangGraph？", "LangGraph 用显式状态图表达节点、条件边、循环、并行、interrupt 和 checkpoint，适合长运行、有回滚和人工审批的 Agent。它比隐藏在 while 循环中的控制流更可观察。", "AgentGraphState 和 retrieval_context、planner、approval_gate、executor、reviewer、rollback、finish 构成主图。", "框架不会自动解决 Prompt、工具安全和业务一致性。"),
        q("StateGraph 中 State 的作用是什么？", "State 是节点间共享的数据契约，节点读取已有字段并返回增量更新。良好 State 应类型明确、字段可序列化、避免把运行时对象和大文件无限塞入。", "AgentGraphState 保存任务、上下文、决策、proposal、测试、review、branch_history、snapshot 和 next_node。", "checkpoint 时特别要关注不可序列化对象和敏感信息。"),
        q("节点和条件边如何设计？", "节点应该对应职责清晰、可重试的步骤；条件边根据结构化状态选择下一节点。业务分支放在图上，底层安全校验留在工具代码。", "planner 根据 action 路由 retrieval_context、approval_gate、executor 或 reviewer；reviewer 根据 approve/revise/reject 路由 finish 或 rollback。", "过细的节点增加复杂度，过粗的节点失去可观察性。"),
        q("如何实现循环和最大轮数？", "rollback 或 read_more 可以返回 Planner，State 中保存 turn；Planner 每轮递增，超过 max_turns 后进入 finish/exhausted，防止无限循环。", "_planner_node 默认 max_turns=4，递归限制还设置为 64。", "轮数、token、工具调用和总耗时都可以作为预算。"),
        q("LangGraph 的 Send 有什么用途？", "Send 支持从一个节点动态分发多个并行分支，每个分支接收不同输入，结果通过 reducer 聚合。适合并行检索、候选规划和独立评审。", "ParallelBranchGraphState 的 candidates 使用 Annotated[list, add] 聚合，dispatch 为五种 profile 创建 Send。", "并行会增加模型调用和限流压力，需要控制 fan-out。"),
        q("什么是 Checkpointer？", "Checkpointer 保存每个 thread 的图状态，使任务可恢复、回放、审计，并支持 interrupt 后继续。thread_id 用于隔离不同任务。", "run_agent_graph 接受 checkpointer、thread_id 和 resume；有 checkpointer 时 configurable.thread_id 默认使用 task_id。", "当前是否配置持久化 checkpointer要如实回答；接口存在不代表已在生产启用。"),
        q("interrupt 和 resume 如何工作？", "节点调用 interrupt 暂停图并返回待人工处理的数据；外部保存状态并收集用户响应，再用 Command(resume=...) 继续同一个 thread。", "approval_gate 返回候选路径、operation 和 summary；resume=true 或 approved=true 后进入 executor。", "没有持久化 checkpointer 时，跨进程恢复不可靠。"),
        q("如何设计 LangGraph fallback？", "依赖缺失或框架不可用时可以执行功能等价的原生路径，但必须保持相同安全工具和状态语义，并记录 fallback。", "run_subgoal_graph、run_task_lifecycle_graph 和 parallel branch 都有 ImportError fallback；聊天捕获 LangGraphUnavailable。", "回退不能静默改变权限或跳过 Reviewer。"),
        q("如何测试 LangGraph？", "节点测试验证输入输出，路由测试覆盖每个条件边，端到端测试覆盖 approve、revise、reject、interrupt、恢复和最大轮数。外部 LLM 使用假响应。", "test_api.py 验证 Agent 事件、patch、review、rollback、memory、artifacts 和聊天追问；还应增加独立图状态测试。", "测试最终状态之外，还要断言没有越权副作用。"),
        q("LangChain 和 LangGraph 如何分工？", "LangChain 更偏模型、Prompt、Retriever、Tool 和结构化输出组件；LangGraph更偏有状态编排。两者可以组合，也可以只使用 LangGraph 加自定义模型工具层。", "code_helper 目前主要使用 LangGraph 编排，自定义 urllib LLM 客户端和 assistant_tools，没有强依赖完整 LangChain。", "不要为了框架统一而重写成熟的安全工具。"),
    ]),
    ("六、记忆、RAG 与状态管理", [
        q("短期记忆和长期记忆有什么区别？", "短期记忆是当前任务状态和最近对话，生命周期较短；长期记忆跨任务保存总结、经验和相关文件，后续通过检索注入上下文。", "AgentGraphState/current_context 是短期状态；SQLite task_memory 保存 run_summary、task_plan、lesson 等长期记忆。", "长期记忆必须可检索、可更新、可删除，并处理错误记忆。"),
        q("RAG 在 Agent 中起什么作用？", "RAG 先检索相关知识，再把证据交给模型，降低凭空生成并控制上下文规模。Agent 可根据任务动态发起多轮检索。", "search_workspace_files 根据标题、描述、focus_paths 和 memory hints 检索代码，read_more 会回到 retrieval_context。", "代码检索还可以升级为 AST、符号索引、向量检索和混合排序。"),
        q("如何设计记忆写入？", "不要保存所有原始对话，应保存结构化、高价值、可复用的总结，并记录来源、时间和相关文件。写入前可以让 Reviewer 或规则检查。", "任务完成后 TaskService 生成 run_summary、task_plan、lesson_success 等 MemoryRecord，包含 keywords 和 related_files。", "失败经验也有价值，但要避免把偶发错误写成永久规则。"),
        q("如何检索记忆？", "使用任务文本提取关键词，结合标题、内容、标签、文件和时间做评分；规模大时使用向量数据库和重排序。", "SQLiteStore.search_memory 对 title、content、keywords、related_files 做加权，并加入时间衰减。", "当前实现不是向量检索，面试中应如实说明。"),
        q("如何避免记忆污染？", "保存时标记来源和置信度，区分成功/失败，允许过期和删除；检索后仍把记忆当参考证据，而不是最高优先级指令。", "memory hints 只辅助 Planner，文件内容、测试和 Reviewer 仍是事实依据。", "敏感源码和用户信息还要考虑隔离与保留周期。"),
        q("Agent 状态应该存在哪里？", "短任务可在内存中，长任务应持久化数据库或 LangGraph checkpointer；大文件放对象存储，State 只保存引用和摘要。", "任务、事件和结果存 SQLite；Agent 运行态在 LangGraph State；checkpointer 接口已预留但是否持久化启用要看运行配置。", "多实例场景必须让状态和锁离开单进程内存。"),
        q("事件和状态有什么区别？", "状态描述当前值，事件描述发生过什么。状态便于查询，事件便于审计、重放和前端实时更新。可靠系统通常同时维护状态表和追加事件。", "tasks 表保存 status/current_step，task_events 按 sequence 保存 planner、patch、test、review 等事件，前端通过 SSE 消费。", "要处理事件重复、乱序和状态更新失败。"),
    ]),
    ("七、可靠性、评测与可观测性", [
        q("如何评测一个 Agent？", "至少分任务成功率、工具调用正确率、事实性、安全性、成本、延迟和用户满意度。评测集要包含正常、边界、对抗和失败恢复场景。", "可以用测试任务验证文件是否真实修改、pytest 是否通过、Reviewer 是否批准、最终回答是否覆盖要求。", "只看模型输出好不好不够，必须评测环境副作用。"),
        q("什么是离线评测和在线评测？", "离线评测使用固定任务集可重复比较版本；在线评测观察真实流量的成功率、回滚率、延迟、成本和人工反馈。二者结合才能持续优化。", "test_api.py 可作为离线回归起点；task_events 和 branch_comparison 可以支持线上分析。", "线上实验要设置安全阈值和快速回滚。"),
        q("如何构建 Agent 测试集？", "从真实任务抽取代表性样例，标注预期文件、允许操作、测试结果和回答要求；加入模糊需求、恶意 Prompt、空仓库、测试失败和工具超时。", "项目已有修复 formatter、创建 sorter、聊天追问、命令隔离和篡改快照测试，可扩展为 Agent benchmark。", "不要只用模型容易完成的 happy path。"),
        q("如何减少幻觉？", "提供检索证据、要求结构化引用实际路径、让工具返回真实结果、Reviewer 交叉检查，并禁止无证据声称成功。", "Chat Prompt 明确要求只有 matching events、patch.applied、changed_files 和 test results 才能声称完成。", "幻觉无法完全消除，应通过架构把影响限制在可恢复范围。"),
        q("如何处理模型服务超时或限流？", "设置客户端超时、指数退避、最大重试、熔断和降级；保存任务状态以便恢复。不要无限重试同一请求。", "LLMConfig 有 timeout_seconds，LLM 不可用时 Planner、聊天和图有 heuristic/fallback 路径。", "降级结果必须向用户说明，尤其不能把未使用 LLM 的结果伪装成模型审查。"),
        q("如何控制 Agent 成本？", "限制最大轮数、候选数量、上下文长度、模型等级和工具调用；缓存稳定结果；仅在高价值节点使用强模型。", "max_turns 限制循环，文件上下文截断，parallel branch 固定五个候选；这些都可进一步配置为预算。", "并行多 Agent 是主要成本来源，应通过评测决定是否保留。"),
        q("需要记录哪些可观测数据？", "记录 task/thread ID、节点、模型、token、延迟、工具参数摘要、结果、重试、错误、路由、测试和最终决策。敏感内容需脱敏。", "事件系统记录 sequence、type、step、payload 和 duration；branch_comparison 保存候选和选择结果。", "可接入 OpenTelemetry，将 FastAPI 请求、LangGraph 节点和 LLM 调用串成 trace。"),
        q("如何避免 Agent 卡死？", "设置节点和工具超时、最大递归和轮数、重复动作检测、取消信号、预算和 watchdog；任务卡住时保存状态并允许人工终止。", "图 config 设置 recursion_limit，Planner 有 max_turns，pytest 有 timeout；前端还显示当前状态。", "需要区分真实长测试和状态未更新，使用心跳与耗时指标判断。"),
    ]),
    ("八、code_helper 项目表达", [
        q("请用两分钟介绍 Agent 架构。", "系统分为 Chat Agent 和代码任务 Agent。Chat Agent 负责意图分类与回复审查；任务 Agent 使用 LangGraph 编排检索、Planner、审批、Executor、Reviewer、Rollback、Finish。LLM 负责语义判断，工具层负责安全执行，SQLite 和事件系统负责持久化与展示。", "关键文件是 llm.py、langgraph_workflow.py、parallel_branches.py、pipeline.py、task_service.py 和 assistant_tools。", "最后主动说明当前是本地 MVP，持久化 checkpointer、多实例队列和权限系统仍可加强。"),
        q("一次代码任务如何流转？", "任务创建后加载 memory 和 workspace context，Planner 并行生成候选并选择动作；缺少信息则 read_more，有补丁则审批/执行；Executor 快照、写入和测试；Reviewer 批准或回滚重试；Finish 汇总并持久化。", "LangGraph 条件边直接表达这些路径，TaskService 外层还有 prepare→execute→persist 生命周期图。", "面试时把 API 层、Agent 图和工具层分开讲。"),
        q("为什么同时有多个 LangGraph？", "不同层级的图负责不同生命周期：任务主图负责 Agent 循环，子目标图负责 inspect/implement/verify，聊天图负责意图路由，生命周期图负责 prepare/execute/persist，并行图负责候选 fan-out。", "这种分层减少单个巨型图的复杂度，每个子图拥有更窄的 State。", "也要防止过度拆图，跨图状态和错误传播必须清晰。"),
        q("如何保证 Planner 生成完整补丁？", "Prompt 要求 create 提供完整文件，modify 提供能在上下文精确匹配的 old 和完整 replacement；解析和工具层再次校验。无效 proposal 不进入写入。", "_resolve_react_proposal 和 PatchProposal 验证路径、operation、old/new；llm.plan.invalid 会触发重新规划。", "长文件可能超过输出限制，需要分块编辑或更强的编辑工具协议。"),
        q("为什么 Reviewer 不能只看测试通过？", "测试可能不完整、没有覆盖用户要求，甚至没有收集到测试。Reviewer 还要检查文件是否真实变化、补丁范围、任务交付项和安全性。", "review_result 同时读取 latest_test、changed_files、patch、branch_history 和 task_description；失败测试不能 approve。", "测试是证据之一，不是唯一真相。"),
        q("如何处理没有测试的项目？", "明确区分“测试通过”和“没有测试”。可以运行语法检查、目标程序或静态检查，但必须在最终结果中说明验证范围。", "run_pytest 对没有任何测试文件且 return_code=5 的情况标记 verification skipped，而不是伪造具体用例通过。", "是否接受无测试应由任务类型和 Reviewer 规则决定。"),
        q("为什么前端只显示聚合过程？", "内部事件很多，全部显示会淹没用户。后端保留完整事件用于审计，前端只聚合关键子目标和耗时，终态展示一问一答。", "TaskDetail 根据 goal.started/completed、test.completed 和 task terminal events 构建过程卡片；Timeline 仍能查看原始事件。", "隐藏的是内部噪声，不是删除审计数据，也不展示隐藏思维链。"),
        q("多轮追问如何触发新任务？", "Chat Agent 根据完整上下文判断 implementation intent；如果需要改代码，TaskService 组合原始任务与新指令，等待当前任务结束或直接启动 follow-up。", "_pending_follow_up_tasks 防止追问与活动执行相互覆盖，聊天事件保持消息顺序。", "多 Worker 时进程内队列需要迁移到持久化任务系统。"),
        q("记忆如何影响下一次任务？", "任务结束后生成结构化记忆，下次根据标题、关键词和相关文件检索，把高分结果作为 Planner 上下文和 memory hints。", "SQLite task_memory 保存 run_summary、task_plan、lesson；search_memory 加权评分并加载前五条。", "记忆只是辅助，不应覆盖当前仓库事实。"),
        q("当前系统最大的生产化风险是什么？", "主要包括进程内后台任务和 SSE 订阅、多 Worker 不共享状态、SQLite 并发、持久化 checkpointer 未必启用、认证隔离不足、LLM 成本和外部服务稳定性。", "这些问题不影响本地 MVP，但多人和多实例部署前要迁移 PostgreSQL、Redis/队列、共享事件总线和 OIDC。", "能主动指出限制会让项目表达更可信。"),
    ]),
    ("九、生产化与系统设计", [
        q("如何把 Agent 部署为多实例？", "把任务队列、锁、事件和 checkpoint 放到外部系统；每个 Worker 无状态消费任务；数据库记录幂等键和租约；SSE 从共享事件总线读取。", "可将 SQLite 改 PostgreSQL，_running_tasks 改 Redis/Celery，EventService 的内存订阅改 Redis Streams/Kafka。", "还要控制同一 workspace 的并发写入。"),
        q("如何处理同一仓库并发修改？", "使用 workspace 级锁、Git worktree 或临时副本隔离任务；写入前检查版本或文件哈希；冲突时重新规划而不是覆盖。", "当前任务直接操作所选 workspace，适合单任务；规模化时应为 task 创建 worktree，并把基线 commit 写入 State。", "数据库锁不能解决文件系统和 Git 冲突。"),
        q("如何实现任务取消？", "状态中记录 cancel_requested，节点和长工具定期检查；subprocess 支持终止；取消后清理资源、恢复快照并写终态事件。", "当前 TaskStatus 有 cancelled，但完整取消链路仍可加强；pytest subprocess 和 LLM 请求需要可取消封装。", "取消必须幂等，不能留下半写文件。"),
        q("如何支持任务恢复？", "使用持久化 checkpointer 保存每个节点状态，工具操作记录幂等 ID；启动时扫描 running 状态并恢复或标记失败。", "run_agent_graph 已支持 checkpointer/thread_id/resume，TaskService 还需要配置持久化实现并保存审批状态。", "State 中的大对象应改为数据库或产物引用。"),
        q("如何进行模型路由？", "根据任务类型、风险、上下文长度和预算选择模型；简单分类用小模型，复杂代码规划和审查用强模型；失败时升级。", "当前 OpenAIPlanner 使用统一配置，未来可以为 Chat、Planner、Reviewer 分别配置模型和超时。", "路由效果要通过离线评测和成本指标验证。"),
        q("如何保护用户源码隐私？", "明确数据边界、最小化发送内容、脱敏密钥、配置模型数据保留、租户隔离、访问审计和删除策略；高敏代码可使用本地模型。", "当前只截取相关文件片段发送 LLM，但生产版仍应识别 .env、密钥和敏感目录并排除。", "日志、memory 和 artifacts 也可能保存源码，必须纳入治理。"),
        q("如何设计 Agent API？", "API 应异步化：创建任务返回 task_id/202，查询状态和结果，SSE/WebSocket 推进度，支持取消、审批和恢复；请求包含 workspace、目标和安全策略。", "code_helper 已有 /tasks、/run、/chat、/events、/rollback 和 artifacts，人工审批 API 可作为下一步。", "接口要有幂等键、权限和版本控制。"),
        q("如何做灰度发布和回滚？", "Prompt、模型和图版本都要可追踪；新版本先离线评测，再小流量灰度；监控成功率、回滚率、成本和安全告警；异常时切回旧图或旧模型。", "事件 payload 可加入 workflow_version、prompt_version 和 model；当前已有 provider/model/branch 信息，可继续扩展。", "Agent 发布不能只看服务是否启动，要看任务质量是否退化。"),
    ]),
    ("十、现场设计题与回答模板", [
        q("如何现场设计一个代码修复 Agent？", "先定义输入与成功标准，再画状态图：检索→规划→审批→补丁→测试→审查→完成/回滚；列出 State、工具 schema、安全限制、最大轮数、事件和评测。", "可以直接借鉴 AgentGraphState 和 code_helper 主图，但要说明哪些组件可替换。", "先讲边界和失败恢复，再讲模型与 Prompt。"),
        q("如何现场设计一个客服 Agent？", "将知识检索、意图分类、回答、工具操作、人工转接和质检拆成节点；退款等高风险工具需要权限和审批；所有回答基于客户和订单事实。", "可复用 chat graph 的 intent route 和 reviewer 模式，但工具和状态换成订单、用户与工单。", "重点是权限、隐私、审计和人工兜底。"),
        q("如何现场设计一个数据分析 Agent？", "节点包括 schema discovery、查询规划、SQL 生成、只读执行、结果校验、图表和解释；数据库账号只读，SQL 做 AST/白名单校验并限制扫描量。", "与 code_helper 类似，LLM 生成候选动作，确定性工具控制执行，Reviewer 检查结果。", "不能给模型任意数据库写权限。"),
        q("如何解释一次 Agent 失败？", "按状态和事件复盘：输入是否清晰、检索是否命中、Prompt 是否产生合法动作、工具是否失败、Reviewer 是否正确、回滚是否完成。用证据定位节点，不只说模型不稳定。", "task_events、branch_history、execution_error、test outcome 和 snapshots 提供完整排查链路。", "再给出短期修复、回归测试和长期评测改进。"),
        q("如何比较两种 Agent 方案？", "使用相同任务集和预算，对比成功率、工具错误率、平均轮数、延迟、成本、安全事件和人工介入率；不能只比较单个演示。", "可以比较单 Planner 与五分支 Planner、原生循环与 LangGraph、无 Reviewer 与独立 Reviewer。", "统计显著性和任务分层也很重要。"),
        q("如何回答“为什么不用纯规则”？", "规则适合确定边界和已知模式，LLM 适合理解自然语言和模糊上下文。最佳方案是规则保证安全与一致性，LLM补充语义决策。", "code_helper 的任务意图、补丁规划和审查使用 LLM，路径、命令、快照和状态迁移使用代码。", "这体现可控 Agent，而不是完全自主。"),
        q("如何回答“为什么不用一个大模型一次完成”？", "一次调用难以获得真实工具反馈、独立验证和失败恢复，也容易把计划误当结果。分节点允许读取更多上下文、执行测试、审查和回滚。", "Planner→Executor→Reviewer 的闭环确保最终结论基于实际 workspace 和 pytest。", "分步会增加延迟和成本，需要根据任务风险选择。"),
        q("如何回答项目亮点和不足？", "亮点要用可验证设计表达：显式 LangGraph、并行候选、工具安全、快照回滚、事件/SSE、记忆和回复审查。不足要诚实：进程内任务、SQLite、多实例和 checkpoint 配置。", "用“当前设计→为什么→限制→下一步”讲 code_helper，避免只列技术名词。", "面试官通常更看重你是否理解取舍和边界。"),
    ]),
]


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "LLM Agent 面试问答 · code_helper 实战版"
    for run in header.runs:
        set_font(run, 8.5, GRAY)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(95)
    kicker.paragraph_format.space_after = Pt(12)
    set_font(kicker.add_run("LLM AGENT INTERVIEW GUIDE"), 11, BLUE, True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("Agent 面试问答"), 30, NAVY, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(subtitle.add_run("结合 code_helper 的多 Agent、LangGraph、工具与安全实践"), 14, GRAY)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(90)
    set_font(meta.add_run("适用于：LLM 应用开发 / Agent 工程 / AI 后端 / LangGraph 面试"), 10.5, GRAY)
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    shade_cell(callout.cell(0, 0), LIGHT_BLUE)
    p = callout.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("回答方法：先讲 Agent 原理，再给 code_helper 证据，最后说明限制与生产化改进。"), 10.5, NAVY, True)
    doc.add_page_break()

    doc.add_heading("项目 Agent 架构速览", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    for cell, text in zip(table.rows[0].cells, ("层次", "当前实现")):
        shade_cell(cell, LIGHT_BLUE)
        set_font(cell.paragraphs[0].add_run(text), 10, NAVY, True)
    rows = [
        ("聊天层", "Chat Agent 分类 implementation/question/panel/status 等意图，Chat Reviewer 复核回复。"),
        ("任务层", "LangGraph：retrieval_context → planner → approval_gate/executor → reviewer → rollback/finish。"),
        ("并行规划", "Planner、Critic、Memory、Explorer、Verifier 五个候选通过 Send 并行构建并评分选择。"),
        ("工具层", "read_file、write_file、run_tests、finish；真实权限由 assistant_tools 控制。"),
        ("记忆层", "SQLite task_memory 保存总结、计划、经验、关键词和相关文件。"),
        ("可观测层", "task_events、SSE、branch_comparison、test duration、snapshot 和 artifacts。"),
        ("安全层", "工作区路径校验、补丁匹配、命令白名单、快照完整性、超时和 Reviewer。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for index, cell in enumerate(cells):
            for run in cell.paragraphs[0].runs:
                set_font(run, 10, DARK_BLUE if index == 0 else RGBColor(0, 0, 0), index == 0)
    set_table_geometry(table, [2700, 6660])
    doc.add_heading("关键代码路径", level=1)
    for text in [
        "packages/agent-core/src/agent_core/workflow/langgraph_workflow.py：主 Agent 图、聊天图、子目标图、审批和回退。",
        "packages/agent-core/src/agent_core/workflow/parallel_branches.py：并行候选 Planner 和评分选择。",
        "packages/agent-core/src/agent_core/llm.py：Prompt、结构化输出、Planner、Executor Router、Reviewer、Chat Agent。",
        "packages/agent-core/src/agent_core/workflow/pipeline.py：代码检索、子目标、测试、结果模型和图入口。",
        "packages/tools/src/assistant_tools/：文件、快照、回滚、命令白名单和 pytest。",
        "apps/api/services/task_service.py：任务生命周期、聊天追问、记忆和最终持久化。",
    ]:
        body(doc, text)
    doc.add_page_break()

    count = 0
    for section_title, items in SECTIONS:
        doc.add_heading(section_title, level=1)
        for item in items:
            count += 1
            question(doc, count, item)

    doc.add_heading("附录：Agent 面试前检查清单", level=1)
    for text in [
        "能画出 Agent 主图并解释每条条件边。",
        "能区分模型决策、图编排和工具权限。",
        "能解释 Planner、Executor、Reviewer、Chat Agent 的职责。",
        "能说明 ReAct、RAG、长期记忆、checkpointer、interrupt 和 Send。",
        "能设计工具 schema、路径校验、命令白名单和幂等策略。",
        "能讲清 Agent 评测、安全、成本、延迟和故障恢复。",
        "能主动指出 code_helper 的 SQLite、进程内任务、多实例和 checkpoint 限制。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(text), 10.5)
    body(doc, f"本手册共整理 {count} 道 Agent 面试问题，内容以当前 code_helper 实现为依据。", "版本说明：")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
